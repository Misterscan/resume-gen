from google import genai
from google.genai import types

import argparse
import json
import logging
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from docx import Document as create_document
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.styles.style import ParagraphStyle
from fpdf import FPDF
from pydantic import BaseModel, Field, ValidationError

# Google API Imports
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
import io

from prompts import (
    RESUME_SYSTEM_PROMPT,
    REVISION_SYSTEM_PROMPT,
    COVER_LETTER_SYSTEM_PROMPT,
)

SCOPES = ['https://www.googleapis.com/auth/drive.file']

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)
logger = logging.getLogger(__name__)

MODEL_NAME = "gemini-3.1-pro-preview"

# --- Models ---

class WorkExperience(BaseModel):
    title: str = ""
    company: str = ""
    location: str = ""
    dates: str = ""
    bullets: List[str] = Field(default_factory=list)

class EducationItem(BaseModel):
    institution: str = ""
    degree: str = ""
    location: str = ""
    dates: str = ""
    details: Union[str, List[str]] = ""

class ResumeSchema(BaseModel):
    professional_summary: str = ""
    work_experience: List[WorkExperience] = Field(default_factory=list)
    education: List[EducationItem] = Field(default_factory=list)
    skills: List[str] = Field(default_factory=list)

class CandidateSchema(BaseModel):
    full_name: str = ""
    contact_info: Union[str, Dict[str, Any]] = ""

class RevisionResponseSchema(BaseModel):
    candidate: CandidateSchema
    resume: ResumeSchema

class CoverLetterSchema(BaseModel):
    recipient_info: str = ""
    greeting: str = ""
    introduction: str = ""
    body_paragraphs: List[str] = Field(default_factory=list)
    company_connection: str = ""
    closing: str = ""
    sign_off: str = ""

# --- PDF Configuration ---



def require_api_key() -> str:
    try:
        api_key = os.environ["GEMINI_API_KEY"]
    except KeyError:
        logger.error("ERROR: GEMINI_API_KEY environment variable is not set.")
        sys.exit(1)

    if not api_key.strip():
        logger.error("ERROR: GEMINI_API_KEY is empty.")
        sys.exit(1)

    logger.warning("SECURITY NOTE: Ensure your API key is not hardcoded or exposed in cleartext.")

    return api_key

def get_google_credentials() -> Any:
    creds = None
    token_path = 'token.json'
    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not os.path.exists('client_secret.json'):
                logger.error("ERROR: client_secret.json not found for Google Docs integration.")
                sys.exit(1)
            flow = InstalledAppFlow.from_client_secrets_file('client_secret.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_path, 'w') as token:
            token.write(creds.to_json())
    return creds

def get_gdoc_text(file_id: str, creds: Any) -> str:
    try:
        drive_service = build('drive', 'v3', credentials=creds)
        request = drive_service.files().export_media(fileId=file_id, mimeType='text/plain')
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        return fh.getvalue().decode('utf-8')
    except Exception as exc:
        logger.error(f"ERROR: Failed to read Google Doc: {exc}")
        sys.exit(1)

def upload_to_gdoc(file_path: Path, file_name: str, creds: Any, overwrite_id: Optional[str] = None) -> str:
    try:
        drive_service = build('drive', 'v3', credentials=creds)
        media = MediaFileUpload(str(file_path), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
        if overwrite_id:
            uploaded_file = drive_service.files().update(
                fileId=overwrite_id,
                media_body=media,
                fields='id, webViewLink'
            ).execute()
        else:
            file_metadata = {
                'name': file_name.replace('.docx', ''),
                'mimeType': 'application/vnd.google-apps.document'
            }
            uploaded_file = drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, webViewLink'
            ).execute()
        return uploaded_file.get('webViewLink', '')
    except Exception as exc:
        logger.error(f"ERROR: Failed to upload to Google Docs: {exc}")
        sys.exit(1)


def prompt_required(label: str) -> str:
    while True:
        value = input(f"{label}: ").strip()
        if value:
            return value
        print("Required field.")


def prompt_optional(label: str) -> str:
    return input(f"{label}: ").strip()


def prompt_yes_no(label: str, default: str = "n") -> bool:
    valid_defaults = {"y", "n"}
    if default not in valid_defaults:
        default = "n"

    suffix = "[Y/n]" if default == "y" else "[y/N]"
    while True:
        value = input(f"{label} {suffix}: ").strip().lower()
        if not value:
            value = default
        if value in {"y", "yes"}:
            return True
        if value in {"n", "no"}:
            return False
        print("Enter y or n.")


def collect_bullets() -> List[str]:
    bullets = []
    print("Enter bullets/responsibilities. Leave blank when done.")
    while True:
        bullet = input("Bullet: ").strip()
        if not bullet:
            break
        bullets.append(bullet)
    return bullets


def collect_revision_notes() -> str:
    print("\nRevision Data")
    print("Enter missing data, corrections, target role, metrics, keywords, or notes.")
    print("Leave blank when done.")

    notes = []
    while True:
        note = input("Revision note: ").strip()
        if not note:
            break
        notes.append(note)

    return "\n".join(notes)


def extract_docx_text(input_path: Path) -> str:
    try:
        document = create_document(str(input_path))
        paragraphs = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        text = "\n".join(paragraphs).strip()
        if not text:
            raise ValueError("The DOCX file did not contain readable text.")
        return text
    except Exception as exc:
        logger.error(f"ERROR: Failed to read DOCX file: {exc}")
        sys.exit(1)


def get_user_input() -> Dict[str, Any]:
    print("Resume Data Collection")

    data: Dict[str, Any] = {
        "full_name": prompt_required("Full Name"),
        "contact_info": prompt_required("Contact Info (Email, Phone, LinkedIn)"),
        "summary": prompt_required(
            "Summary (1-2 sentences highlighting your experience and goals)"
        ),
        "work_experience": [],
        "education": [],
        "skills": [],
    }

    print("\nWork Experience")
    while True:
        add_item = input("Add work experience? [y/n]: ").strip().lower()
        if add_item != "y":
            break

        data["work_experience"].append(
            {
                "title": prompt_required("Title"),
                "company": prompt_required("Company"),
                "location": prompt_required("Location"),
                "dates": prompt_required("Dates"),
                "bullets": collect_bullets(),
            }
        )

    print("\nEducation")
    while True:
        add_item = input("Add education? [y/n]: ").strip().lower()
        if add_item != "y":
            break

        data["education"].append(
            {
                "institution": prompt_required("Institution"),
                "degree": prompt_required("Degree"),
                "dates": prompt_optional("Dates"),
                "location": prompt_required("Location"),
                "details": prompt_optional("Details"),
            }
        )

    print("\nSkills")
    skills_raw = prompt_required("Skills, comma-separated")
    data["skills"] = [
        skill.strip() for skill in skills_raw.split(",") if skill.strip()
    ]

    return data


def call_gemini_json(
    api_key: str,
    system_prompt: str,
    payload: Dict[str, Any],
    temperature: float = 0.3,
) -> Dict[str, Any]:
    try:
        client = genai.Client(api_key=api_key)

        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=json.dumps(payload, indent=2),
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                temperature=temperature,
                thinking_config=types.ThinkingConfig(
                    thinking_level=types.ThinkingLevel.LOW
                ),
                response_mime_type="application/json",
            ),
        )

        response_text = response.text
        if not response_text:
            raise ValueError("Gemini returned an empty response.")
            
        import re
        match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if match:
            response_text = match.group(0)

        try:
            parsed = json.loads(response_text)
            if not isinstance(parsed, dict):
                raise ValueError("Gemini returned JSON, but not a JSON object.")
        except json.JSONDecodeError as exc:
            logger.error(f"ERROR: Failed to parse JSON from Gemini. Raw response:\n{response_text}")
            raise exc

        return parsed

    except Exception as exc:
        logger.error(f"ERROR: Gemini generation failed: {exc}")
        sys.exit(1)


def generate_resume_content(
    raw_data: Dict[str, Any],
    api_key: str,
) -> Dict[str, Any]:
    parsed = call_gemini_json(
        api_key=api_key,
        system_prompt=RESUME_SYSTEM_PROMPT,
        payload={"raw_data": raw_data},
        temperature=0.3,
    )
    try:
        validated = ResumeSchema.model_validate(parsed)
        return validated.model_dump()
    except ValidationError as e:
        logger.error(f"ERROR: Resume schema validation failed:\n{e}")
        sys.exit(1)



def revise_resume_content(
    api_key: str,
    revision_notes: str,
    raw_data: Optional[Dict[str, Any]] = None,
    current_resume: Optional[Dict[str, Any]] = None,
    resume_text: Optional[str] = None,
) -> Dict[str, Any]:
    
    # Token efficiency step: truncate extremely long work experience lists
    # We only send the 5 most recent roles to the LLM to avoid token exhaustion.
    truncated_resume = None
    if current_resume:
        truncated_resume = dict(current_resume)
        if "work_experience" in truncated_resume and isinstance(truncated_resume["work_experience"], list):
            if len(truncated_resume["work_experience"]) > 5:
                logger.info(f"Truncating current resume from {len(truncated_resume['work_experience'])} to 5 recent roles to save tokens.")
                truncated_resume["work_experience"] = truncated_resume["work_experience"][:5]

    payload: Dict[str, Any] = {
        "revision_notes": revision_notes,
        "raw_data": raw_data or {},
        "current_resume": truncated_resume or {},
        "existing_resume_text": resume_text or "",
    }

    result = call_gemini_json(
        api_key=api_key,
        system_prompt=REVISION_SYSTEM_PROMPT,
        payload=payload,
        temperature=0.2,
    )

    try:
        validated = RevisionResponseSchema.model_validate(result)
        out = validated.model_dump()
        # Coerce dictionary contact_info back to string if AI hallucinates JSON object
        if isinstance(out.get('candidate', {}).get('contact_info'), dict):
            out['candidate']['contact_info'] = " | ".join(str(v) for v in out['candidate']['contact_info'].values())
        return out
    except ValidationError as e:
        logger.error(f"ERROR: Revision schema validation failed:\n{e}")
        sys.exit(1)


def generate_cover_letter_content(
    raw_data: Dict[str, Any],
    resume: Dict[str, Any],
    revision_notes: Optional[str],
    target_role: str,
    target_company: str,
    api_key: str,
) -> Dict[str, Any]:
    parsed = call_gemini_json(
        api_key=api_key,
        system_prompt=COVER_LETTER_SYSTEM_PROMPT,
        payload={
            "raw_data": raw_data,
            "resume": resume,
            "revision_notes": revision_notes or "",
            "target_role": target_role,
            "target_company": target_company,
        },
        temperature=0.4,
    )
    
    try:
        validated = CoverLetterSchema.model_validate(parsed)
        return validated.model_dump()
    except ValidationError as e:
        logger.error(f"ERROR: Cover letter schema validation failed:\n{e}")
        sys.exit(1)


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)


def configure_docx(document: Document) -> None:
    normal_style = document.styles["Normal"]

    if normal_style.type == WD_STYLE_TYPE.PARAGRAPH:
        paragraph_style = normal_style
        assert isinstance(paragraph_style, ParagraphStyle)
        paragraph_style.font.name = "Arial"
        paragraph_style.font.size = Pt(10)


def save_as_docx(
    resume: Dict[str, Any],
    raw_data: Dict[str, Any],
    output_path: Path,
) -> None:
    try:
        document = create_document()
        configure_docx(document)

        name = document.add_paragraph()
        name_run = name.add_run(raw_data["full_name"])
        name_run.bold = True
        name_run.font.size = Pt(16)

        document.add_paragraph(raw_data["contact_info"])

        add_heading(document, "Professional Summary")
        document.add_paragraph(resume.get("professional_summary", ""))

        add_heading(document, "Work Experience")
        for item in resume.get("work_experience", []):
            role = document.add_paragraph()
            role_run = role.add_run(
                f"{item.get('title', '')} | "
                f"{item.get('company', '')} | "
                f"{item.get('location', '')} | "
                f"{item.get('dates', '')}"
            )
            role_run.bold = True

            for bullet in item.get("bullets", []):
                document.add_paragraph(bullet, style="List Bullet")

        add_heading(document, "Education")
        for item in resume.get("education", []):
            edu = document.add_paragraph()
            edu_run = edu.add_run(
                f"{item.get('degree', '')} | "
                f"{item.get('institution', '')} | "
                f"{item.get('location', '')} | "
                f"{item.get('dates', '')}"
            )
            edu_run.bold = True

            details = item.get("details", "")
            if isinstance(details, list):
                details = ", ".join(details)
            if details:
                document.add_paragraph(details)

        add_heading(document, "Skills")
        document.add_paragraph(", ".join(resume.get("skills", [])))

        document.save(str(output_path))

    except Exception as exc:
        logger.error(f"ERROR: Failed to write DOCX: {exc}")
        sys.exit(1)


def save_cover_letter_as_docx(
    cover_letter: Dict[str, Any],
    raw_data: Dict[str, Any],
    output_path: Path,
) -> None:
    try:
        document = create_document()
        configure_docx(document)

        name = document.add_paragraph()
        name_run = name.add_run(raw_data["full_name"])
        name_run.bold = True
        name_run.font.size = Pt(16)

        document.add_paragraph(raw_data["contact_info"])
        document.add_paragraph("")  # spacing

        # Date
        import datetime
        document.add_paragraph(datetime.date.today().strftime("%B %d, %Y"))
        document.add_paragraph("")  # spacing

        if cover_letter.get("recipient_info"):
            document.add_paragraph(cover_letter["recipient_info"])
            document.add_paragraph("")

        if cover_letter.get("greeting"):
            document.add_paragraph(cover_letter["greeting"])
            document.add_paragraph("")

        if cover_letter.get("introduction"):
            document.add_paragraph(cover_letter["introduction"])
            document.add_paragraph("")

        for paragraph in cover_letter.get("body_paragraphs", []):
            document.add_paragraph(paragraph)
            document.add_paragraph("")

        if cover_letter.get("company_connection"):
            document.add_paragraph(cover_letter["company_connection"])
            document.add_paragraph("")

        if cover_letter.get("closing"):
            document.add_paragraph(cover_letter["closing"])
            document.add_paragraph("")

        if cover_letter.get("sign_off"):
            document.add_paragraph(cover_letter["sign_off"])

        document.add_paragraph(raw_data.get("full_name", ""))

        document.save(str(output_path))

    except Exception as exc:
        logger.error(f"ERROR: Failed to write Cover Letter DOCX: {exc}")
        sys.exit(1)


class PDFConfig:
    MARGIN = 15.0
    FONT_FAMILY = "Helvetica"
    FONT_SIZE_HEADING = 11
    FONT_SIZE_BODY = 10
    FONT_SIZE_TITLE = 16
    LINE_HEIGHT_NORMAL = 5
    LINE_HEIGHT_HEADING = 7
    LINE_HEIGHT_TITLE = 8
    SPACING_LARGE = 5
    SPACING_MEDIUM = 4
    SPACING_SMALL = 2
    SPACING_TINY = 1


class ResumePDF(FPDF):
    def usable_width(self) -> float:
        return self.w - self.l_margin - self.r_margin

    @staticmethod
    def clean_text(text: Any) -> str:
        return (
            str(text)
            .replace("\u2013", "-")
            .replace("\u2014", "-")
            .replace("\u2018", "'")
            .replace("\u2019", "'")
            .replace("\u201c", '"')
            .replace("\u201d", '"')
            .strip()
        )

    def section_heading(self, text: str) -> None:
        self.ln(PDFConfig.SPACING_MEDIUM)
        self.set_font(PDFConfig.FONT_FAMILY, "B", PDFConfig.FONT_SIZE_HEADING)
        self.set_x(self.l_margin)
        self.cell(
            self.usable_width(),
            PDFConfig.LINE_HEIGHT_HEADING,
            self.clean_text(text).upper(),
            new_x="LMARGIN",
            new_y="NEXT",
        )
        self.set_font(PDFConfig.FONT_FAMILY, "", PDFConfig.FONT_SIZE_BODY)

    def paragraph(self, text: Any) -> None:
        self.set_x(self.l_margin)
        self.multi_cell(
            self.usable_width(),
            PDFConfig.LINE_HEIGHT_NORMAL,
            self.clean_text(text),
            new_x="LMARGIN",
            new_y="NEXT",
        )
        self.ln(PDFConfig.SPACING_TINY)

    def bullet(self, text: Any) -> None:
        self.set_x(self.l_margin)
        self.multi_cell(
            self.usable_width(),
            PDFConfig.LINE_HEIGHT_NORMAL,
            f"- {self.clean_text(text)}",
            new_x="LMARGIN",
            new_y="NEXT",
        )

    def bold_line(self, text: Any) -> None:
        self.set_font(PDFConfig.FONT_FAMILY, "B", PDFConfig.FONT_SIZE_BODY)
        self.set_x(self.l_margin)
        self.multi_cell(
            self.usable_width(),
            PDFConfig.LINE_HEIGHT_NORMAL,
            self.clean_text(text),
            new_x="LMARGIN",
            new_y="NEXT",
        )
        self.set_font(PDFConfig.FONT_FAMILY, "", PDFConfig.FONT_SIZE_BODY)


def create_pdf() -> ResumePDF:
    pdf = ResumePDF()
    pdf.set_left_margin(PDFConfig.MARGIN)
    pdf.set_right_margin(PDFConfig.MARGIN)
    pdf.set_auto_page_break(auto=True, margin=PDFConfig.MARGIN)
    pdf.add_page()
    return pdf


def write_pdf_header(pdf: ResumePDF, raw_data: Dict[str, Any]) -> None:
    pdf.set_font(PDFConfig.FONT_FAMILY, "B", PDFConfig.FONT_SIZE_TITLE)
    pdf.set_x(pdf.l_margin)
    pdf.cell(
        pdf.usable_width(),
        PDFConfig.LINE_HEIGHT_TITLE,
        pdf.clean_text(raw_data["full_name"]),
        new_x="LMARGIN",
        new_y="NEXT",
    )

    pdf.set_font(PDFConfig.FONT_FAMILY, "", PDFConfig.FONT_SIZE_BODY)
    pdf.paragraph(raw_data.get("contact_info", ""))


def save_as_pdf(
    resume: Dict[str, Any],
    raw_data: Dict[str, Any],
    output_path: Path,
) -> None:
    try:
        pdf = create_pdf()
        write_pdf_header(pdf, raw_data)

        pdf.section_heading("Professional Summary")
        pdf.paragraph(resume.get("professional_summary", ""))

        pdf.section_heading("Work Experience")
        for item in resume.get("work_experience", []):
            pdf.bold_line(
                f"{item.get('title', '')} | "
                f"{item.get('company', '')} | "
                f"{item.get('location', '')} | "
                f"{item.get('dates', '')}"
            )

            for bullet in item.get("bullets", []):
                pdf.bullet(bullet)

            pdf.ln(PDFConfig.SPACING_SMALL)

        pdf.section_heading("Education")
        for item in resume.get("education", []):
            pdf.bold_line(
                f"{item.get('degree', '')} | "
                f"{item.get('institution', '')} | "
                f"{item.get('location', '')} | "
                f"{item.get('dates', '')}"
            )

            details = item.get("details", "")
            if isinstance(details, list):
                details = ", ".join(details)
            if details:
                pdf.paragraph(details)

        pdf.section_heading("Skills")
        pdf.paragraph(", ".join(resume.get("skills", [])))

        pdf.output(str(output_path))

    except Exception as exc:
        logger.error(f"ERROR: Failed to write PDF: {exc}")
        sys.exit(1)


def save_cover_letter_as_pdf(
    cover_letter: Dict[str, Any],
    raw_data: Dict[str, Any],
    output_path: Path,
) -> None:
    try:
        pdf = create_pdf()
        write_pdf_header(pdf, raw_data)

        pdf.ln(PDFConfig.SPACING_LARGE)

        import datetime
        pdf.paragraph(datetime.date.today().strftime("%B %d, %Y"))
        pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("recipient_info"):
            for line in cover_letter["recipient_info"].split("\n"):
                pdf.paragraph(line)
            pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("greeting"):
            pdf.paragraph(cover_letter["greeting"])
            pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("introduction"):
            pdf.paragraph(cover_letter["introduction"])
            pdf.ln(PDFConfig.SPACING_SMALL)

        for paragraph in cover_letter.get("body_paragraphs", []):
            pdf.paragraph(paragraph)
            pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("company_connection"):
            pdf.paragraph(cover_letter["company_connection"])
            pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("closing"):
            pdf.paragraph(cover_letter["closing"])
            pdf.ln(PDFConfig.SPACING_MEDIUM)

        if cover_letter.get("sign_off"):
            pdf.paragraph(cover_letter["sign_off"])

        pdf.paragraph(raw_data.get("full_name", ""))

        pdf.output(str(output_path))

    except Exception as exc:
        logger.error(f"ERROR: Failed to write Cover Letter PDF: {exc}")
        sys.exit(1)

def get_base_path(output_arg: Optional[str], full_name: str, custom_dir: Optional[str] = None) -> Path:
    if output_arg:
        target = output_arg
    elif full_name:
        target = f"{full_name} Resume"
    else:
        target = "My Resume"
    
    target = target.replace("/", "_").replace("\\", "_")
    output_path = Path(target).expanduser().resolve()
    
    if custom_dir:
        output_folder = Path(custom_dir).expanduser().resolve()
    else:
        output_folder = output_path.parent / "resumes"
        
    output_folder.mkdir(parents=True, exist_ok=True)
    return output_folder / output_path.stem

def save_resume_files(
    resume: Dict[str, Any],
    raw_data: Dict[str, Any],
    base_path: Path,
    output_format: str,
    args: Optional[argparse.Namespace] = None,
) -> None:
    if output_format in {"docx", "both", "gdocs"}:
        docx_path = base_path.with_suffix(".docx")
        save_as_docx(resume, raw_data, docx_path)
        
        if output_format == "gdocs" and args:
            creds = get_google_credentials()
            overwrite_id = args.gdoc_id if getattr(args, "gdoc_update", False) else None
            link = upload_to_gdoc(docx_path, docx_path.name, creds, overwrite_id)
            logger.info(f"Resume Google Doc saved: {link}")
            try:
                os.remove(docx_path)
            except OSError:
                pass


    if output_format in {"pdf", "both"}:
        save_as_pdf(resume, raw_data, base_path.with_suffix(".pdf"))


def save_cover_letter_files(
    cover_letter: Dict[str, Any],
    raw_data: Dict[str, Any],
    base_path: Path,
    output_format: str,
    args: Optional[argparse.Namespace] = None,
) -> None:
    cl_base_path = base_path.with_name(f"{base_path.name} Cover Letter")

    if output_format in {"docx", "both", "gdocs"}:
        docx_path = cl_base_path.with_suffix(".docx")
        save_cover_letter_as_docx(cover_letter, raw_data, docx_path)
        
        if output_format == "gdocs" and args:
            creds = get_google_credentials()
            link = upload_to_gdoc(docx_path, docx_path.name, creds, None) # cover letter never overwrites existing resume
            logger.info(f"Cover Letter Google Doc saved: {link}")
            try:
                os.remove(docx_path)
            except OSError:
                pass


    if output_format in {"pdf", "both"}:
        save_cover_letter_as_pdf(cover_letter, raw_data, cl_base_path.with_suffix(".pdf"))


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate an ATS-optimized resume and optional cover letter with Gemini."
    )
    parser.add_argument(
        "--format",
        choices=["pdf", "docx", "both", "gdocs"],
        default="both",
        help="Output format. 'gdocs' will create a Google Doc in your Drive.",
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Target base filename without extension. Defaults to '<Full Name> Resume'.",
    )
    parser.add_argument(
        "--dir",
        default=None,
        help="Target output directory. Defaults to './resumes'.",
    )
    parser.add_argument(
        "--revise",
        action="store_true",
        help="Prompt for missing data or corrections after initial resume generation.",
    )
    parser.add_argument(
        "--cl",
        action="store_true",
        help="Automatically generate a tailored cover letter after resume creation or revision.",
    )
    parser.add_argument(
        "--input",
        type=Path,
        help="Existing .docx resume to revise instead of starting from scratch.",
    )
    parser.add_argument(
        "--gdoc-id",
        help="Existing Google Doc ID to revise instead of a local file.",
    )
    parser.add_argument(
        "--gdoc-update",
        action="store_true",
        help="Update the target Google Doc when using --gdoc-id instead of creating a new file.",
    )
    parser.add_argument(
        "--notes",
        help="Revision instructions for --input mode.",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Print generated JSON to the console instead of saving files.",
    )
    return parser.parse_args()



def setup_environment(output_dir: Optional[str]) -> str:
    """Verifies API key and ensures the target output directory exists."""
    api_key = require_api_key()
    
    if output_dir:
        out_path = Path(output_dir).expanduser().resolve()
    else:
        out_path = Path.cwd() / "resumes"
        
    out_path.mkdir(parents=True, exist_ok=True)
    logger.info(f"Target output directory verified: {out_path}")
    
    return api_key

def print_diff_summary(old: Dict[str, Any], new: Dict[str, Any]) -> None:
    logger.info("--- Resume Revision Diff Summary ---")
    if old.get("professional_summary") != new.get("professional_summary"):
        logger.info("* Professional Summary was updated.")
        
    old_exp = old.get("work_experience", [])
    new_exp = new.get("work_experience", [])
    if len(old_exp) != len(new_exp):
        logger.info(f"* Work experience count changed: {len(old_exp)} -> {len(new_exp)}")
    else:
        for i, (o, n) in enumerate(zip(old_exp, new_exp)):
            if o != n:
                logger.info(f"* Work experience '{n.get('title')} at {n.get('company')}' was updated.")
                
    old_edu = old.get("education", [])
    new_edu = new.get("education", [])
    if old_edu != new_edu:
        logger.info("* Education section was updated.")
        
    old_skills = set(old.get("skills", []))
    new_skills = set(new.get("skills", []))
    if old_skills != new_skills:
        added = new_skills - old_skills
        removed = old_skills - new_skills
        if added: logger.info(f"* Skills Added: {', '.join(added)}")
        if removed: logger.info(f"* Skills Removed: {', '.join(removed)}")
    logger.info("------------------------------------")


def main() -> None:
    args = parse_args()
    api_key = setup_environment(args.dir)

    if args.input or args.gdoc_id:
        if args.gdoc_id:
            creds = get_google_credentials()
            resume_text = get_gdoc_text(args.gdoc_id, creds)
            
            revision_notes = args.notes or collect_revision_notes()
            if not revision_notes:
                logger.error("ERROR: Revision notes are required for revision.")
                sys.exit(1)
        else:
            input_path = args.input.expanduser().resolve()
            if input_path.suffix.lower() != ".docx":
                logger.error("ERROR: --input currently supports .docx files only.")
                sys.exit(1)
            if not input_path.exists():
                logger.error(f"ERROR: Input file not found: {input_path}")
                sys.exit(1)

            resume_text = extract_docx_text(input_path)
            revision_notes = args.notes or collect_revision_notes()
            if not revision_notes:
                logger.error("ERROR: Revision notes are required in --input mode.")
                sys.exit(1)

        revised = revise_resume_content(
            api_key=api_key,
            revision_notes=revision_notes,
            resume_text=resume_text,
        )
        raw_data = revised["candidate"]
        resume = revised["resume"]
        raw_data["source_resume_text"] = resume_text

        if args.dry_run:
            logger.info("DRY RUN: Outputting Revised Resume JSON:")
            print(json.dumps(resume, indent=2))
        else:
            base_path = get_base_path(args.output, raw_data.get("full_name", ""), args.dir)
            save_resume_files(resume, raw_data, base_path, args.format, args)

        should_generate_cl = args.cl or prompt_yes_no(
            "Generate a tailored Cover Letter from the revised resume?",
            default="n",
        )
        if should_generate_cl:
            target_role = prompt_optional("Target Role / Job Title")
            target_company = prompt_optional("Target Company")
            cover_letter = generate_cover_letter_content(
                raw_data=raw_data,
                resume=resume,
                revision_notes=revision_notes,
                target_role=target_role,
                target_company=target_company,
                api_key=api_key,
            )
            if args.dry_run:
                logger.info("DRY RUN: Outputting Cover Letter JSON:")
                print(json.dumps(cover_letter, indent=2))
            else:
                save_cover_letter_files(cover_letter, raw_data, base_path, args.format, args)

        print("Generation complete.")
        return

    raw_data = get_user_input()
    resume = generate_resume_content(raw_data, api_key)

    revision_notes = ""
    should_revise = args.revise or prompt_yes_no(
        "Revise resume with missing data or corrections?",
        default="n",
    )

    if should_revise:
        revision_notes = collect_revision_notes()
        if revision_notes:
            revised = revise_resume_content(
                api_key=api_key,
                revision_notes=revision_notes,
                raw_data=raw_data,
                current_resume=resume,
            )
            candidate = revised["candidate"]
            raw_data["full_name"] = candidate.get("full_name", raw_data.get("full_name", ""))
            raw_data["contact_info"] = candidate.get("contact_info", raw_data.get("contact_info", ""))
            
            print_diff_summary(resume, revised["resume"])
            resume = revised["resume"]
        else:
            logger.info("No revision notes entered. Keeping initial resume.")

    if args.dry_run:
        logger.info("DRY RUN: Outputting Final Resume JSON:")
        print(json.dumps(resume, indent=2))
    else:
        base_path = get_base_path(args.output, raw_data.get("full_name", ""), args.dir)
        save_resume_files(resume, raw_data, base_path, args.format, args)

    should_generate_cl = args.cl or prompt_yes_no(
        "Generate a tailored Cover Letter from the final resume?",
        default="n",
    )

    if should_generate_cl:
        target_role = prompt_optional("Target Role / Job Title")
        target_company = prompt_optional("Target Company")
        cover_letter = generate_cover_letter_content(
            raw_data=raw_data,
            resume=resume,
            revision_notes=revision_notes,
            target_role=target_role,
            target_company=target_company,
            api_key=api_key,
        )
        if args.dry_run:
            logger.info("DRY RUN: Outputting Cover Letter JSON:")
            print(json.dumps(cover_letter, indent=2))
        else:
            save_cover_letter_files(cover_letter, raw_data, base_path, args.format, args)

    print("Generation complete.")


if __name__ == "__main__":
    main()
