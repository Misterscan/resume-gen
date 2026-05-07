import logging
import io
from typing import Any, Dict
from docx import Document as create_document
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from fpdf import FPDF

from .exceptions import DocumentError

logger = logging.getLogger(__name__)

MAX_EXTRACTED_DOCX_CHARS = 80_000

def _iter_table_text(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                value = paragraph.text.strip()
                if value:
                    yield value


def extract_docx_text(input_file: Any) -> str:
    try:
        document = create_document(input_file)
        parts: list[str] = []

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                value = paragraph.text.strip()
                if value:
                    parts.append(value)

        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if value:
                parts.append(value)

        for table in document.tables:
            parts.extend(_iter_table_text(table))

        for section in document.sections:
            for paragraph in section.footer.paragraphs:
                value = paragraph.text.strip()
                if value:
                    parts.append(value)

        text = "\n".join(parts).strip()

        if not text:
            raise DocumentError("The DOCX file did not contain readable text.")

        if len(text) > MAX_EXTRACTED_DOCX_CHARS:
            raise DocumentError(
                f"The DOCX text is too large: {len(text)} characters."
            )

        return text

    except DocumentError:
        raise
    except Exception as exc:
        logger.info("DOCX extraction failed: %s", exc.__class__.__name__)
        raise DocumentError("Failed to read DOCX content.") from exc

def candidate_name(raw_data: Dict[str, Any]) -> str:
    value = str(raw_data.get("full_name") or "").strip()
    return value or "Candidate"


def candidate_contact(raw_data: Dict[str, Any]) -> str:
    return str(raw_data.get("contact_info") or "").strip()


def format_role_line(item: Dict[str, Any]) -> str:
    parts = [
        item.get("title", ""),
        item.get("company", ""),
        item.get("location", ""),
        item.get("dates", ""),
    ]
    return " | ".join(str(part).strip() for part in parts if str(part).strip())


def format_education_line(item: Dict[str, Any]) -> str:
    parts = [
        item.get("degree", ""),
        item.get("institution", ""),
        item.get("location", ""),
        item.get("dates", ""),
    ]
    return " | ".join(str(part).strip() for part in parts if str(part).strip())


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)

def configure_docx(document: Document) -> None:
    normal_style = document.styles["Normal"]
    # Provide a type hint/cast for the linter since it might not know paragraph styles
    if normal_style.type == WD_STYLE_TYPE.PARAGRAPH:
        paragraph_style: Any = normal_style
        paragraph_style.font.name = "Arial"
        paragraph_style.font.size = Pt(10)

def generate_docx_stream(resume: Dict[str, Any], raw_data: Dict[str, Any]) -> io.BytesIO:
    try:
        document = create_document()
        configure_docx(document)

        name = document.add_paragraph()
        name_run = name.add_run(candidate_name(raw_data))
        name_run.bold = True
        name_run.font.size = Pt(16)

        document.add_paragraph(candidate_contact(raw_data))

        add_heading(document, "Professional Summary")
        document.add_paragraph(resume.get("professional_summary", ""))

        add_heading(document, "Work Experience")
        for item in resume.get("work_experience", []):
            role = document.add_paragraph()
            role_run = role.add_run(format_role_line(item))
            role_run.bold = True
            for bullet in item.get("bullets", []):
                document.add_paragraph(bullet, style="List Bullet")

        add_heading(document, "Education")
        for item in resume.get("education", []):
            edu = document.add_paragraph()
            edu_run = edu.add_run(format_education_line(item))
            edu_run.bold = True
            details = item.get("details", "")
            if isinstance(details, list):
                details = ", ".join(details)
            if details:
                document.add_paragraph(details)

        add_heading(document, "Skills")
        document.add_paragraph(", ".join(resume.get("skills", [])))

        stream = io.BytesIO()
        document.save(stream)
        stream.seek(0)
        return stream
    except Exception as exc:
        logger.error(f"Failed to write DOCX: {exc}")
        raise DocumentError(f"Failed to generate DOCX: {exc}")

def generate_cover_letter_docx_stream(cover_letter: Dict[str, Any], raw_data: Dict[str, Any]) -> io.BytesIO:
    try:
        document = create_document()
        configure_docx(document)

        name = document.add_paragraph()
        name_run = name.add_run(candidate_name(raw_data))
        name_run.bold = True
        name_run.font.size = Pt(16)

        document.add_paragraph(candidate_contact(raw_data))
        document.add_paragraph("")

        import datetime
        document.add_paragraph(datetime.date.today().strftime("%B %d, %Y"))
        document.add_paragraph("")

        if cover_letter.get("recipient_info"):
            document.add_paragraph(cover_letter["recipient_info"])
            document.add_paragraph("")

        if cover_letter.get("greeting"):
            document.add_paragraph(cover_letter["greeting"])
            document.add_paragraph("")

        if cover_letter.get("introduction"):
            document.add_paragraph(cover_letter["introduction"])
            document.add_paragraph("")

        for paragraph in cover_letter.get("body_paragraphs", []):
            document.add_paragraph(paragraph)
            document.add_paragraph("")

        if cover_letter.get("company_connection"):
            document.add_paragraph(cover_letter["company_connection"])
            document.add_paragraph("")

        if cover_letter.get("closing"):
            document.add_paragraph(cover_letter["closing"])
            document.add_paragraph("")

        if cover_letter.get("sign_off"):
            document.add_paragraph(cover_letter["sign_off"])

        document.add_paragraph(raw_data.get("full_name", ""))

        stream = io.BytesIO()
        document.save(stream)
        stream.seek(0)
        return stream
    except Exception as exc:
        logger.error(f"Failed to write Cover Letter DOCX: {exc}")
        raise DocumentError(f"Failed to generate Cover Letter DOCX: {exc}")


class PDFConfig:
    MARGIN = 15.0
    FONT_FAMILY = "Helvetica"
    FONT_SIZE_HEADING = 11
    FONT_SIZE_BODY = 10
    FONT_SIZE_TITLE = 16
    LINE_HEIGHT_NORMAL = 5
    LINE_HEIGHT_HEADING = 7
    LINE_HEIGHT_TITLE = 8
    SPACING_LARGE = 5
    SPACING_MEDIUM = 4
    SPACING_SMALL = 2
    SPACING_TINY = 1

class ResumePDF(FPDF):
    def usable_width(self) -> float:
        return self.w - self.l_margin - self.r_margin

    @staticmethod
    def clean_text(text: Any) -> str:
        return (
            str(text)
            .replace("\u2013", "-")
            .replace("\u2014", "-")
            .replace("\u2018", "'")
            .replace("\u2019", "'")
            .replace("\u201c", '"')
            .replace("\u201d", '"')
            .strip()
        )

    def section_heading(self, text: str) -> None:
        self.ln(PDFConfig.SPACING_MEDIUM)
        self.set_font(PDFConfig.FONT_FAMILY, "B", PDFConfig.FONT_SIZE_HEADING)
        self.set_x(self.l_margin)
        self.cell(
            self.usable_width(),
            PDFConfig.LINE_HEIGHT_HEADING,
            self.clean_text(text).upper(),
            new_x="LMARGIN",
            new_y="NEXT",
        )
        self.set_font(PDFConfig.FONT_FAMILY, "", PDFConfig.FONT_SIZE_BODY)

    def paragraph(self, text: Any) -> None:
        self.set_x(self.l_margin)
        self.multi_cell(
            self.usable_width(),
            PDFConfig.LINE_HEIGHT_NORMAL,
            self.clean_text(text),
            new_x="LMARGIN",
            new_y="NEXT",
        )
        self.ln(PDFConfig.SPACING_TINY)

    def bullet(self, text: Any) -> None:
        self.set_x(self.l_margin)
        self.multi_cell(
            self.usable_width(),
            PDFConfig.LINE_HEIGHT_NORMAL,
            f"- {self.clean_text(text)}",
            new_x="LMARGIN",
            new_y="NEXT",
        )

    def bold_line(self, text: Any) -> None:
        self.set_font(PDFConfig.FONT_FAMILY, "B", PDFConfig.FONT_SIZE_BODY)
        self.set_x(self.l_margin)
        self.multi_cell(
            self.usable_width(),
            PDFConfig.LINE_HEIGHT_NORMAL,
            self.clean_text(text),
            new_x="LMARGIN",
            new_y="NEXT",
        )
        self.set_font(PDFConfig.FONT_FAMILY, "", PDFConfig.FONT_SIZE_BODY)

def create_pdf() -> ResumePDF:
    pdf = ResumePDF()
    pdf.set_left_margin(PDFConfig.MARGIN)
    pdf.set_right_margin(PDFConfig.MARGIN)
    pdf.set_auto_page_break(auto=True, margin=PDFConfig.MARGIN)
    pdf.add_page()
    return pdf

def write_pdf_header(pdf: ResumePDF, raw_data: Dict[str, Any]) -> None:
    pdf.set_font(PDFConfig.FONT_FAMILY, "B", PDFConfig.FONT_SIZE_TITLE)
    pdf.set_x(pdf.l_margin)
    pdf.cell(
        pdf.usable_width(),
        PDFConfig.LINE_HEIGHT_TITLE,
        pdf.clean_text(candidate_name(raw_data)),
        new_x="LMARGIN",
        new_y="NEXT",
    )
    pdf.set_font(PDFConfig.FONT_FAMILY, "", PDFConfig.FONT_SIZE_BODY)
    contact = candidate_contact(raw_data)
    if contact:
        pdf.paragraph(contact)

def _finalize_pdf_stream(pdf: FPDF) -> io.BytesIO:
    stream = io.BytesIO()
    raw_pdf = pdf.output()
    # Handle both binary and string outputs depending on FPDF version
    if isinstance(raw_pdf, str):
        stream.write(raw_pdf.encode('latin1'))
    else:
        stream.write(raw_pdf)
    stream.seek(0)
    return stream

def generate_pdf_stream(resume: Dict[str, Any], raw_data: Dict[str, Any]) -> io.BytesIO:
    try:
        pdf = create_pdf()
        write_pdf_header(pdf, raw_data)

        pdf.section_heading("Professional Summary")
        pdf.paragraph(resume.get("professional_summary", ""))

        pdf.section_heading("Work Experience")
        for item in resume.get("work_experience", []):
            pdf.bold_line(
                f"{item.get('title', '')} | "
                f"{item.get('company', '')} | "
                f"{item.get('location', '')} | "
                f"{item.get('dates', '')}"
            )
            for bullet in item.get("bullets", []):
                pdf.bullet(bullet)
            pdf.ln(PDFConfig.SPACING_SMALL)

        pdf.section_heading("Education")
        for item in resume.get("education", []):
            pdf.bold_line(
                f"{item.get('degree', '')} | "
                f"{item.get('institution', '')} | "
                f"{item.get('location', '')} | "
                f"{item.get('dates', '')}"
            )
            details = item.get("details", "")
            if isinstance(details, list):
                details = ", ".join(details)
            if details:
                pdf.paragraph(details)

        pdf.section_heading("Skills")
        pdf.paragraph(", ".join(resume.get("skills", [])))

        return _finalize_pdf_stream(pdf)
    except Exception as exc:
        logger.error(f"Failed to write PDF: {exc}")
        raise DocumentError(f"Failed to generate PDF: {exc}")

def generate_cover_letter_pdf_stream(cover_letter: Dict[str, Any], raw_data: Dict[str, Any]) -> io.BytesIO:
    try:
        pdf = create_pdf()
        write_pdf_header(pdf, raw_data)
        pdf.ln(PDFConfig.SPACING_LARGE)

        import datetime
        pdf.paragraph(datetime.date.today().strftime("%B %d, %Y"))
        pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("recipient_info"):
            for line in cover_letter["recipient_info"].split("\n"):
                pdf.paragraph(line)
            pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("greeting"):
            pdf.paragraph(cover_letter["greeting"])
            pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("introduction"):
            pdf.paragraph(cover_letter["introduction"])
            pdf.ln(PDFConfig.SPACING_SMALL)

        for paragraph in cover_letter.get("body_paragraphs", []):
            pdf.paragraph(paragraph)
            pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("company_connection"):
            pdf.paragraph(cover_letter["company_connection"])
            pdf.ln(PDFConfig.SPACING_SMALL)

        if cover_letter.get("closing"):
            pdf.paragraph(cover_letter["closing"])
            pdf.ln(PDFConfig.SPACING_MEDIUM)

        if cover_letter.get("sign_off"):
            pdf.paragraph(cover_letter["sign_off"])

        pdf.paragraph(raw_data.get("full_name", ""))

        return _finalize_pdf_stream(pdf)
    except Exception as exc:
        logger.error(f"Failed to write Cover Letter PDF: {exc}")
        raise DocumentError(f"Failed to generate Cover Letter PDF: {exc}")

