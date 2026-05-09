import io
import pytest
from docx import Document

from services.document import (
    extract_docx_text,
    generate_docx_stream,
    generate_pdf_stream,
    generate_cover_letter_docx_stream,
)
from services.exceptions import DocumentError


@pytest.fixture
def resume_data():
    return {
        "professional_summary": "Experienced software engineer.",
        "work_experience": [
            {
                "title": "Senior Engineer",
                "company": "Tech Corp",
                "location": "Remote",
                "dates": "Jan 2020 - Present",
                "bullets": ["Optimized backend.", "Led team."],
            }
        ],
        "projects": [
            {
                "name": "Capstone Project",
                "organization": "University",
                "location": "Remote",
                "dates": "2024",
                "bullets": ["Built a full-stack app."],
            }
        ],
        "volunteer_experience": [
            {
                "role": "Volunteer Tutor",
                "organization": "Community Center",
                "location": "City, State",
                "dates": "2023 - Present",
                "bullets": ["Tutored students weekly."],
            }
        ],
        "certifications": [
            {
                "name": "Google Data Analytics Certificate",
                "issuer": "Google",
                "dates": "2025",
                "details": "Completed coursework in SQL and spreadsheets.",
            }
        ],
        "education": [
            {
                "degree": "B.S. Computer Science",
                "institution": "University",
                "location": "City, State",
                "dates": "2015 - 2019",
                "details": "GPA: 4.0",
            }
        ],
        "skills": ["Python", "Django", "TDD"],
    }


@pytest.fixture
def raw_data():
    return {"full_name": "Test User", "contact_info": "test@example.com | 555-5555"}


@pytest.fixture
def cover_letter_data():
    return {
        "recipient_info": "Hiring Manager, Tech Corp",
        "greeting": "Dear Hiring Manager,",
        "introduction": "I am writing to apply for the position.",
        "body_paragraphs": ["I have 5 years of Python experience."],
        "company_connection": "I love Tech Corp's mission.",
        "closing": "I look forward to an interview.",
        "sign_off": "Sincerely,",
    }


def _docx_text(stream):
    doc = Document(stream)
    return "\n".join(p.text for p in doc.paragraphs)


def test_generate_docx_stream_returns_valid_bytesio(resume_data, raw_data):
    result_stream = generate_docx_stream(resume_data, raw_data)

    assert isinstance(result_stream, io.BytesIO)

    text = _docx_text(result_stream)
    assert "Test User" in text
    assert "test@example.com | 555-5555" in text
    assert "PROFESSIONAL SUMMARY" in text
    assert "Experienced software engineer" in text
    assert "PROJECTS" in text
    assert "VOLUNTEER EXPERIENCE" in text
    assert "CERTIFICATIONS" in text


def test_generate_pdf_stream_returns_valid_bytesio(resume_data, raw_data):
    result_stream = generate_pdf_stream(resume_data, raw_data)

    assert isinstance(result_stream, io.BytesIO)
    assert len(result_stream.getvalue()) > 0


def test_generate_cover_letter_docx_stream_returns_valid_bytesio(cover_letter_data, raw_data):
    result_stream = generate_cover_letter_docx_stream(cover_letter_data, raw_data)

    assert isinstance(result_stream, io.BytesIO)

    text = _docx_text(result_stream)
    assert "Test User" in text
    assert "test@example.com | 555-5555" in text
    assert "Dear Hiring Manager," in text
    assert "Sincerely," in text


def test_extract_docx_text_valid_file():
    temp_doc = Document()
    temp_doc.add_paragraph("First paragraph text.")
    temp_doc.add_paragraph("Second paragraph text.")
    stream = io.BytesIO()
    temp_doc.save(stream)
    stream.seek(0)

    extracted_text = extract_docx_text(stream)
    assert extracted_text == "First paragraph text.\nSecond paragraph text."


def test_extract_docx_text_empty_file_raises_error():
    temp_doc = Document()
    stream = io.BytesIO()
    temp_doc.save(stream)
    stream.seek(0)

    with pytest.raises(DocumentError):
        extract_docx_text(stream)
